"""
Build a standard patient/family care study Word document (.docx) from study
JSON, matching the Nursing & Midwifery Council of Ghana format used by the
sample care studies in data/templates/ (title page, chapters, numbered
sections, drug/care-plan tables).

Usage:
    python src/export_docx.py < study.json > care_study.docx

Reads a JSON object on stdin with the shape:

    {
      "title": {
        "patientName": "MRS. P.A",
        "diagnosis": "SICKLE CELL-ACUTE CHEST SYNDROME",
        "studentName": "DRAMANI ZENABU",
        "indexNumber": "...",
        "collegeName": "NURSING AND MIDWIFERY TRAINING COLLEGE",
        "collegeLocation": "SEIKWA",
        "year": "2023"
      },
      "chapters": [
        {
          "name": "Assessment",
          "intro": "This chapter opens with... (optional, rendered under the heading)",
          "introReferences": [{"label": "...", "inText": "(WHO, 2024)", "url": "..."}],
          "sections": [
            {
              "id": "1.1",
              "heading": "Patient's Particulars",
              "draft": "**1.1 Patient's Particulars**  The patient ...",
              "references": [
                {"label": "World Health Organization. (2024). Pneumonia in children. Fact sheet.", "inText": "(WHO, 2024)", "url": "https://www.who.int/news-room/fact-sheets/detail/pneumonia"}
              ],
              "fields": [{"label": "Age", "value": "10 years"}],
              "rows": {
                "title": "Prescribed drugs",
                "columns": ["Drug", "Class"],
                "data": [["Ceftriaxone", "3rd-gen cephalosporin"]]
              }
            }
          ]
        }
      ]
    }

Draft text may contain light markdown (**bold**, *italic*, <sup>superscript</sup>,
~~strikethrough~~, ==highlight==, ++underline++, "- " bullets and "1. " numbered
lines — list lines indented by 2+ spaces become nested levels), which is
converted to proper Word formatting.

A paragraph may also carry a "<!-- align:center spacing:1.5 -->" directive line
before it (as emitted by the preview editor's paragraph tools) to override the
document theme's alignment / line spacing for just that paragraph.

The optional "scope" key controls how much of the study is rendered:
    {"type": "full"}                            (default) title page + TOC + all chapters
    {"type": "chapter", "chapterIndex": 0}      one chapter (heading + its sections)
    {"type": "section", "chapterIndex": 0, "sectionIndex": 1}
                                                  a single section only
"""
import argparse
import io
import json
import re
import sys
from dataclasses import dataclass, fields

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROMAN = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]

# ---------------------------------------------------------------------------
# Document theme — every formatting decision flows through this one object so
# the exporter can be restyled without touching render code. The defaults
# reproduce the classic NMC Ghana care-study look exactly. The client may pass
# a {"theme": {...}} block in the payload to override individual knobs.
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class Theme:
    body_font: str = "Times New Roman"
    heading_font: str = "Times New Roman"
    body_size: int = 12
    heading1_size: int = 14  # chapter headings, TABLE OF CONTENTS, REFERENCES
    heading2_size: int = 12  # section headings ("1.1 Patient's Particulars")
    table_size: int = 10
    table_title_size: int = 11
    title_size: int = 14  # patient name / diagnosis on the title page
    body_color: str = "000000"
    heading_color: str = "000000"
    table_header_fill: str = "D9D9D9"
    table_header_color: str = "000000"
    highlight_color: str = "FFFF00"  # ==highlight== runs
    first_line_indent: float = 0.0  # inches; e.g. 0.5 gives the classic essay look
    line_spacing: float = 1.5
    space_after: int = 6
    heading1_space_before: int = 14
    heading1_space_after: int = 8
    heading2_space_before: int = 10
    heading2_space_after: int = 4
    body_alignment: str = "justify"  # justify | left | center | right

    @classmethod
    def from_dict(cls, data):
        """Build a Theme from an optional payload dict; unknown keys are ignored."""
        if not data:
            return cls()
        known = {f.name for f in fields(cls)}
        return cls(**{k: v for k, v in data.items() if k in known})


# Named styles the document is rendered through. Students can restyle the whole
# document from Word's style pane (or by tweaking the theme above).
STYLE_BODY = "Care Study Body"
STYLE_HEADING_1 = "Care Study Heading 1"
STYLE_HEADING_2 = "Care Study Heading 2"
STYLE_TABLE_HEADER = "Care Study Table Header"
STYLE_TABLE_CELL = "Care Study Table Cell"
STYLE_TOC_TITLE = "Care Study TOC Title"


def _alignment_enum(value):
    """Map an alignment name ("left" / "center" / "right" / "justify") to the
    Word paragraph-alignment enum; anything unknown falls back to justify."""
    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get((value or "justify").lower(), WD_ALIGN_PARAGRAPH.JUSTIFY)


def _theme_alignment(theme):
    return _alignment_enum(theme.body_alignment)

# Inline markdown tokens we translate into Word runs.
INLINE_RE = re.compile(
    r"(<sup>[^<]*</sup>|==[^=]+?==|\+\+[^+]+?\+{2}|~~[^~]+?~~|\*\*[^*]+\*\*|\*[^*]+\*)"
)
BULLET_RE = re.compile(r"^\s*[-•]\s+")
NUMBER_RE = re.compile(r"^\s*\d+[.)]\s+")
# Paragraph-style directives the preview editor emits: a "<!-- ... -->" line
# before a paragraph overrides the document theme's alignment / line spacing
# for just that paragraph ("<!-- align:center spacing:1.5 -->").
PARA_DIRECTIVE_RE = re.compile(
    r"^<!--\s*(?:align:(left|center|right|justify))?\s*(?:spacing:(\d+(?:\.\d+)?))?\s*-->\s*$"
)
# Markdown pipe tables that can sneak into drafts (e.g. "| Item | Detail |").
TABLE_ROW_RE = re.compile(r"^\s*\|")
SEPARATOR_RE = re.compile(r"^:?-{2,}:?$")
# Bare day + month (e.g. "1 August") and ISO ("2012-06-11") -> ordinal with superscript.
MONTHS = (
    r"January|February|March|April|May|June|July|August|"
    r"September|October|November|December"
)
BARE_DATE_RE = re.compile(rf"\b(\d{{1,2}})\s+({MONTHS})\b")
ISO_DATE_RE = re.compile(r"\b(\d{4})-(\d{1,2})-(\d{1,2})\b")
MONTH_NAMES = {
    1: "January", 2: "February", 3: "March", 4: "April", 5: "May", 6: "June",
    7: "July", 8: "August", 9: "September", 10: "October", 11: "November", 12: "December",
}


def _apply_font(style, font_name, size_pt, bold=None, color=None):
    """Font a style (name, size, optional bold + color) including the complex-
    script font slots so Word never falls back to its default theme font."""
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def _set_outline_level(style, level):
    """Give a style an outline level so Word's Navigation pane and TOC fields
    pick its paragraphs up (0 = Heading 1, 1 = Heading 2, ...)."""
    ppr = style.element.get_or_add_pPr()
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(level))
    ppr.append(outline)


def _setup_styles(doc, theme):
    """Build the named Care Study styles from the theme. Every paragraph and
    table cell renders through one of these styles, so formatting lives in
    Word's style pane (editable there) rather than on individual runs."""
    normal = doc.styles["Normal"]
    _apply_font(normal, theme.body_font, theme.body_size, color=theme.body_color)
    normal.paragraph_format.line_spacing = theme.line_spacing
    normal.paragraph_format.space_after = Pt(theme.space_after)

    body = doc.styles.add_style(STYLE_BODY, WD_STYLE_TYPE.PARAGRAPH)
    body.base_style = normal
    body.paragraph_format.line_spacing = theme.line_spacing
    body.paragraph_format.space_after = Pt(theme.space_after)
    body.paragraph_format.alignment = _theme_alignment(theme)
    if theme.first_line_indent:
        body.paragraph_format.first_line_indent = Inches(theme.first_line_indent)

    heading1 = doc.styles.add_style(STYLE_HEADING_1, WD_STYLE_TYPE.PARAGRAPH)
    heading1.base_style = normal
    _apply_font(
        heading1, theme.heading_font, theme.heading1_size,
        bold=True, color=theme.heading_color,
    )
    heading1.paragraph_format.space_before = Pt(theme.heading1_space_before)
    heading1.paragraph_format.space_after = Pt(theme.heading1_space_after)
    heading1.paragraph_format.keep_with_next = True

    heading2 = doc.styles.add_style(STYLE_HEADING_2, WD_STYLE_TYPE.PARAGRAPH)
    heading2.base_style = normal
    _apply_font(
        heading2, theme.heading_font, theme.heading2_size,
        bold=True, color=theme.heading_color,
    )
    heading2.paragraph_format.space_before = Pt(theme.heading2_space_before)
    heading2.paragraph_format.space_after = Pt(theme.heading2_space_after)
    heading2.paragraph_format.keep_with_next = True

    # Outline levels put chapters/sections in Word's Navigation pane and feed
    # the auto-updating TOC field (TOC \o "1-3").
    _set_outline_level(heading1, 0)
    _set_outline_level(heading2, 1)

    table_header = doc.styles.add_style(STYLE_TABLE_HEADER, WD_STYLE_TYPE.PARAGRAPH)
    table_header.base_style = normal
    _apply_font(
        table_header, theme.body_font, theme.table_size,
        bold=True, color=theme.table_header_color,
    )

    table_cell = doc.styles.add_style(STYLE_TABLE_CELL, WD_STYLE_TYPE.PARAGRAPH)
    table_cell.base_style = normal
    _apply_font(table_cell, theme.body_font, theme.table_size)

    # Centered page-title style (TABLE OF CONTENTS / REFERENCES) — deliberately
    # no outline level so it never shows in the Navigation pane or TOC field.
    toc_title = doc.styles.add_style(STYLE_TOC_TITLE, WD_STYLE_TYPE.PARAGRAPH)
    toc_title.base_style = normal
    _apply_font(
        toc_title, theme.heading_font, theme.heading1_size,
        bold=True, color=theme.heading_color,
    )
    toc_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Built-in list styles sometimes fall back to the default theme font;
    # pin them to the body font so bullets match the rest of the document.
    for name in ("List Bullet", "List Number"):
        try:
            _apply_font(doc.styles[name], theme.body_font, theme.body_size)
        except KeyError:
            pass


def _add_run(paragraph, text, theme, bold=False, italic=False, size=None, superscript=False,
             underline=False, strike=False, highlight=None, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if strike:
        run.font.strike = True
    if superscript:
        run.font.superscript = True
    if size is not None:
        run.font.size = Pt(size)
    if highlight:
        # True hex highlight via run shading (w:highlight only knows 16 names).
        rpr = run._r.get_or_add_rPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), highlight)
        rpr.append(shd)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = theme.body_font
    return run


def _ordinal_suffix(day):
    """English ordinal suffix for a day number: 1->st, 2->nd, 3->rd, else th."""
    if 10 <= day % 100 <= 20:
        return "th"
    return {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")


def _format_ordinal_dates(text):
    """Normalize dates to an ordinal with a superscript suffix.

    - Bare '1 August 2026'  -> '1<sup>st</sup> August 2026'
    - ISO '2012-06-11'      -> '11<sup>th</sup> June 2012' (date-picker values)

    A safety net for dates the model wrote without an ordinal suffix; dates
    already written as '1<sup>st</sup> August' are left untouched."""

    def iso_repl(match):
        year, month, day = match.group(1), int(match.group(2)), int(match.group(3))
        month_name = MONTH_NAMES.get(month)
        if month_name is None:
            return match.group(0)
        return f"{day}<sup>{_ordinal_suffix(day)}</sup> {month_name} {year}"

    def repl(match):
        day = int(match.group(1))
        return f"{day}<sup>{_ordinal_suffix(day)}</sup> {match.group(2)}"

    return ISO_DATE_RE.sub(iso_repl, BARE_DATE_RE.sub(repl, text))


def _add_inline_text(paragraph, text, theme):
    """Split **bold**, *italic*, <sup>superscript</sup>, ~~strikethrough~~,
    ==highlight== and ++underline++ into styled runs."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("<sup>") and part.endswith("</sup>"):
            _add_run(paragraph, part[5:-6], theme, superscript=True)
        elif part.startswith("==") and part.endswith("=="):
            _add_run(paragraph, part[2:-2], theme, highlight=theme.highlight_color)
        elif part.startswith("++") and part.endswith("++"):
            _add_run(paragraph, part[2:-2], theme, underline=True)
        elif part.startswith("~~") and part.endswith("~~"):
            _add_run(paragraph, part[2:-2], theme, strike=True)
        elif part.startswith("**") and part.endswith("**"):
            _add_run(paragraph, part[2:-2], theme, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            _add_run(paragraph, part[1:-1], theme, italic=True)
        else:
            _add_run(paragraph, part, theme)


def _add_markdown_paragraph(doc, text, theme, style=None, alignment=None, spacing=None):
    """Add a paragraph, honouring **bold**, *italic* and <sup>superscript</sup>.
    Style defaults to Care Study Body; alignment and line spacing default to
    the theme's, unless a paragraph directive overrides them."""
    paragraph = doc.add_paragraph(style=style or STYLE_BODY)
    # alignment may be a WD_PARAGRAPH_ALIGNMENT enum (from callers like the
    # title page) or an "align:..." directive string — accept both.
    if alignment is None:
        paragraph.alignment = _theme_alignment(theme)
    elif isinstance(alignment, str):
        paragraph.alignment = _alignment_enum(alignment)
    else:
        paragraph.alignment = alignment
    paragraph.paragraph_format.line_spacing = (
        spacing if spacing is not None else theme.line_spacing
    )
    _add_inline_text(paragraph, _format_ordinal_dates(text), theme)
    return paragraph





def _split_table_row(line):
    """Split a markdown row like '| A | B |' into trimmed cells."""
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _add_markdown_table(doc, table_lines, theme):
    """Convert a markdown pipe table (header, separator, data rows) to a Word table."""
    rows = []
    for line in table_lines:
        cells = _split_table_row(line)
        # Skip the '|------|--------|' separator row.
        if cells and all(SEPARATOR_RE.match(cell) for cell in cells):
            continue
        rows.append(cells)
    if not rows:
        return

    width = max(len(row) for row in rows)
    header = rows[0] + [""] * (width - len(rows[0]))
    data = [row + [""] * (width - len(row)) for row in rows[1:]]
    _add_data_table(doc, header, data, theme)


def _draft_table_row_count(draft: str) -> int:
    """Count the data rows in the draft's markdown pipe table (0 when absent)."""
    lines = draft.split("\n")
    index = 0
    while index < len(lines):
        if not TABLE_ROW_RE.match(lines[index]):
            index += 1
            continue
        block = []
        while index < len(lines) and TABLE_ROW_RE.match(lines[index]):
            block.append(lines[index])
            index += 1
        count = 0
        for position, line in enumerate(block):
            if position == 0:
                continue  # header row
            cells = _split_table_row(line)
            if cells and all(SEPARATOR_RE.match(cell) for cell in cells):
                continue  # separator row
            count += 1
        return count
    return 0


# ---------------------------------------------------------------------------
# Multi-level list numbering — real Word numPr (numId + ilvl) so bullets and
# numbered lists indent properly and wrapped lines align under the text rather
# than under the marker, instead of flat single-level list styles.
# ---------------------------------------------------------------------------


def _list_level_xml(ilvl, fmt, text, left_twips, rfonts=None):
    """One <w:lvl> entry for the abstract numbering definition."""
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), str(ilvl))
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    numfmt = OxmlElement("w:numFmt")
    numfmt.set(qn("w:val"), fmt)
    lvltext = OxmlElement("w:lvlText")
    lvltext.set(qn("w:val"), text)
    lvljc = OxmlElement("w:lvlJc")
    lvljc.set(qn("w:val"), "left")
    ppr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left_twips))
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    lvl.append(start)
    lvl.append(numfmt)
    lvl.append(lvltext)
    lvl.append(lvljc)
    lvl.append(ppr)
    if rfonts:
        rpr = OxmlElement("w:rPr")
        rf = OxmlElement("w:rFonts")
        rf.set(qn("w:ascii"), rfonts)
        rf.set(qn("w:hAnsi"), rfonts)
        rf.set(qn("w:hint"), "default")
        rpr.append(rf)
        lvl.append(rpr)
    return lvl


def _add_abstract_num(numbering, abstract_id, levels):
    an = OxmlElement("w:abstractNum")
    an.set(qn("w:abstractNumId"), str(abstract_id))
    mlt = OxmlElement("w:multiLevelType")
    mlt.set(qn("w:val"), "hybridMultilevel")
    an.append(mlt)
    for level in levels:
        an.append(level)
    return an


def _add_num(numbering, num_id, abstract_id):
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    return num


def _setup_lists(doc, theme):
    """Create bullet + decimal multi-level numbering definitions (three visual
    levels each) and remember their numIds on the doc for _add_draft."""
    numbering = doc.part.numbering_part.element
    next_abstract = (
        max(
            (int(a.get(qn("w:abstractNumId"))) for a in numbering.findall(qn("w:abstractNum"))),
            default=-1,
        )
        + 1
    )
    next_num = (
        max((int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))), default=0) + 1
    )

    bullet_abstract = _add_abstract_num(
        numbering, next_abstract,
        [
            _list_level_xml(0, "bullet", "\u2022", 720, rfonts="Symbol"),
            _list_level_xml(1, "bullet", "o", 1440, rfonts="Courier New"),
            _list_level_xml(2, "bullet", "\u25aa", 2160),
        ],
    )
    decimal_abstract = _add_abstract_num(
        numbering, next_abstract + 1,
        [
            _list_level_xml(0, "decimal", "%1.", 720),
            _list_level_xml(1, "lowerLetter", "%2.", 1440),
            _list_level_xml(2, "decimal", "%3.", 2160),
        ],
    )
    # Schema: all abstractNum elements precede all num elements.
    first_num = numbering.find(qn("w:num"))
    if first_num is not None:
        first_num.addprevious(bullet_abstract)
        first_num.addprevious(decimal_abstract)
    else:
        numbering.append(bullet_abstract)
        numbering.append(decimal_abstract)
    numbering.append(_add_num(numbering, next_num, next_abstract))
    numbering.append(_add_num(numbering, next_num + 1, next_abstract + 1))

    doc._carestudy_lists = {"bullet": next_num, "decimal": next_num + 1}
    return doc._carestudy_lists


def _apply_list_numbering(paragraph, num_id, level):
    """Point a paragraph at a numbering level (numPr) and match its indent so
    wrapped lines align under the text, not the marker."""
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.get_or_add_numPr()
    numpr.get_or_add_ilvl().set(qn("w:val"), str(level))
    numpr.get_or_add_numId().set(qn("w:val"), str(num_id))
    paragraph.paragraph_format.left_indent = Inches(0.5 * (level + 1))
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)


def _add_draft(doc, draft, theme):
    """Render a drafted section: prose paragraphs, bullets/lists, and any stray
    markdown pipe tables are converted into proper Word tables. "<!-- ... -->"
    directive lines style just the paragraph that follows them."""
    lines = draft.split("\n")
    index = 0
    pending = {}
    while index < len(lines):
        raw = lines[index].rstrip()
        if not raw.strip():
            # A directive separated from its paragraph by a blank line is an
            # orphan; drop it so it never leaks into the document.
            pending = {}
            index += 1
            continue
        directive = PARA_DIRECTIVE_RE.match(raw)
        if directive:
            style = {}
            if directive.group(1):
                style["alignment"] = directive.group(1)
            if directive.group(2):
                style["spacing"] = min(max(float(directive.group(2)), 1.0), 3.0)
            pending = {**pending, **style}  # adjacent directives combine
            index += 1
            continue
        if TABLE_ROW_RE.match(raw):
            # Consume the whole contiguous table block.
            pending = {}
            table_lines = []
            while index < len(lines) and TABLE_ROW_RE.match(lines[index]):
                table_lines.append(lines[index])
                index += 1
            _add_markdown_table(doc, table_lines, theme)
            continue
        bullet = BULLET_RE.match(raw)
        number = NUMBER_RE.match(raw)
        if bullet or number:
            lists = getattr(doc, "_carestudy_lists", None)
            level = min(len(re.match(r"^\s*", raw).group(0)) // 2, 2)
            text = BULLET_RE.sub("", raw) if bullet else NUMBER_RE.sub("", raw)
            paragraph = _add_markdown_paragraph(
                doc, text, theme, style="List Bullet" if bullet else "List Number",
                alignment=pending.pop("alignment", None),
                spacing=pending.pop("spacing", None),
            )
            if lists is not None:
                _apply_list_numbering(
                    paragraph,
                    lists["bullet"] if bullet else lists["decimal"],
                    level,
                )
        else:
            _add_markdown_paragraph(
                doc, raw, theme,
                alignment=pending.pop("alignment", None),
                spacing=pending.pop("spacing", None),
            )
        index += 1


def _shade_cell(cell, fill="D9D9D9"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _add_data_table(doc, header, data, theme):
    """Build a bordered Word table from a header row and data rows, rendered
    through the Care Study styles (bold shaded header via the theme)."""
    if not header:
        return
    width = len(header)
    table = doc.add_table(rows=1 + len(data), cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for col_index, column in enumerate(header):
        cell = table.rows[0].cells[col_index]
        cell.text = ""
        cell.paragraphs[0].style = STYLE_TABLE_HEADER
        _add_inline_text(cell.paragraphs[0], _format_ordinal_dates(column), theme)
        _shade_cell(cell, theme.table_header_fill)

    for row_index, row in enumerate(data, start=1):
        for col_index in range(width):
            cell_value = row[col_index] if col_index < len(row) else ""
            cell = table.rows[row_index].cells[col_index]
            cell.text = ""
            cell.paragraphs[0].style = STYLE_TABLE_CELL
            _add_inline_text(
                cell.paragraphs[0],
                _format_ordinal_dates(cell_value) or "—",
                theme,
            )


def _add_rows_table(doc, rows, theme):
    """Render a repeatable rows section (drugs, care plan, outcomes) as a table."""
    header = doc.add_paragraph(style=STYLE_BODY)
    header.paragraph_format.space_before = Pt(6)
    header.paragraph_format.space_after = Pt(4)
    _add_run(header, rows.get("title", ""), theme, bold=True, size=theme.table_title_size)
    _add_data_table(doc, rows.get("columns") or [], rows.get("data") or [], theme)


def _add_page_number(doc):
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _add_content_control(paragraph, text, alias, tag, theme, bold=False, size=None):
    """Wrap text in a Word content control (fillable field). The alias shows in
    Word's controls UI; the tag is a stable identifier. Formatting (font, bold,
    size) is applied to the run inside the control."""
    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")
    for name, value in (("w:alias", alias), ("w:tag", tag)):
        element = OxmlElement(name)
        element.set(qn("w:val"), value)
        sdt_pr.append(element)
    sdt_content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), theme.body_font)
    rfonts.set(qn("w:hAnsi"), theme.body_font)
    rfonts.set(qn("w:cs"), theme.body_font)
    rpr.append(rfonts)
    if bold:
        rpr.append(OxmlElement("w:b"))
    if size is not None:
        half = str(int(round(size * 2)))
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), half)
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), half)
        rpr.append(sz)
        rpr.append(sz_cs)
    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = text
    run.append(text_el)
    sdt_content.append(run)
    sdt.append(sdt_pr)
    sdt.append(sdt_content)
    paragraph._p.append(sdt)
    return paragraph


def _add_cc_paragraph(doc, text, alias, tag, theme, bold=False, size=None):
    """A centered body paragraph whose text lives in a Word content control."""
    paragraph = doc.add_paragraph(style=STYLE_BODY)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_content_control(paragraph, text, alias, tag, theme, bold=bold, size=size)
    return paragraph


def _add_title_page(doc, title, theme):
    """Centered title page in the style of the sample care studies."""
    center = WD_ALIGN_PARAGRAPH.CENTER
    _add_markdown_paragraph(doc, "PATIENT/FAMILY CARE STUDY", theme, alignment=center).runs[0].bold = True
    _add_markdown_paragraph(doc, "ON", theme, alignment=center)
    patient = (title.get("patientName") or "").strip()
    if patient:
        _add_cc_paragraph(
            doc, patient.upper(), "Patient Name", "PatientName",
            theme, bold=True, size=theme.title_size,
        )
    _add_markdown_paragraph(doc, "WITH", theme, alignment=center)
    diagnosis = (title.get("diagnosis") or "").strip()
    if diagnosis:
        _add_cc_paragraph(
            doc, diagnosis.upper(), "Diagnosis", "Diagnosis",
            theme, bold=True, size=theme.title_size,
        )
    _add_markdown_paragraph(doc, "PRESENTED BY", theme, alignment=center)
    student = (title.get("studentName") or "").strip()
    if student:
        index = (title.get("indexNumber") or "").strip()
        _add_cc_paragraph(
            doc, student.upper() + (f" ({index})" if index else ""),
            "Student Name", "StudentName", theme,
        )
    _add_markdown_paragraph(doc, "A FINAL YEAR STUDENT OF", theme, alignment=center)
    college = (title.get("collegeName") or "").strip()
    location = (title.get("collegeLocation") or "").strip()
    if college:
        _add_cc_paragraph(
            doc, college.upper() + (f", {location.upper()}" if location else ""),
            "College", "CollegeName", theme,
        )
    _add_markdown_paragraph(
        doc,
        "A PATIENT AND FAMILY CARE STUDY SUBMITTED TO NURSING AND MIDWIFERY COUNCIL OF GHANA "
        "IN PARTIAL FULFILLMENT OF THE REQUIREMENT FOR THE AWARD OF LICENSE IN GENERAL NURSING",
        theme,
        alignment=center,
    )
    year = (title.get("year") or "").strip()
    if year:
        p = _add_cc_paragraph(doc, year, "Year", "Year", theme)
        p.paragraph_format.space_before = Pt(24)

    doc.add_page_break()


def _chapter_heading(chapters, chapter_index, chapter):
    """Heading for a chapter: 'CHAPTER I: ASSESSMENT' for numbered chapters, or
    the bare name (e.g. 'PRELIMINARY PAGES') for the unnumbered front matter.
    The ordinal counts only non-front-matter chapters, so adding the
    preliminary pages never shifts the I–VI numbering."""
    if chapter.get("isFrontMatter"):
        return chapter.get("name", "").upper()
    ordinal = sum(
        1 for c in chapters[:chapter_index] if not c.get("isFrontMatter")
    )
    return f"CHAPTER {ROMAN[ordinal]}: {chapter.get('name', '').upper()}"


def _add_toc(doc, chapters, theme):
    """'TABLE OF CONTENTS' followed by a live Word TOC field.

    The field refreshes itself when Word opens the document (or on F9), pulling
    the chapter/section headings through their outline levels. The static
    entries are cached inside the field, so a table of contents is visible even
    before Word recalculates it."""
    heading = doc.add_paragraph(style=STYLE_TOC_TITLE)
    _add_run(heading, "TABLE OF CONTENTS", theme)

    entries = []
    for chapter_index, chapter in enumerate(chapters):
        p = doc.add_paragraph(style=STYLE_BODY)
        p.paragraph_format.space_before = Pt(8)
        _add_run(p, _chapter_heading(chapters, chapter_index, chapter), theme, bold=True)
        entries.append(p)
        for section in chapter.get("sections", []):
            sp = doc.add_paragraph(style=STYLE_BODY)
            sp.paragraph_format.left_indent = Inches(0.4)
            sp.paragraph_format.space_after = Pt(2)
            _add_run(sp, f"{section.get('id', '')} {section.get('heading', '')}", theme)
            entries.append(sp)

    # Field begin/instruction/separate runs lead the first entry paragraph so
    # the cached list shows without a stray blank line above it.
    if entries:
        first, last = entries[0], entries[-1]
    else:
        first = doc.add_paragraph(style=STYLE_BODY)
        last = first

    anchor = first._p.find(qn("w:pPr"))
    field_runs = []
    begin = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    begin.append(fld_begin)
    field_runs.append(begin)

    instr = OxmlElement("w:r")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    instr.append(instr_text)
    field_runs.append(instr)

    separate = OxmlElement("w:r")
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    separate.append(fld_separate)
    field_runs.append(separate)

    for run in field_runs:
        if anchor is not None:
            anchor.addnext(run)
        else:
            first._p.insert(0, run)
        anchor = run

    end = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end.append(fld_end)
    last._p.append(end)

    if not _has_bibliography(chapters):
        references = doc.add_paragraph(style=STYLE_TOC_TITLE)
        references.paragraph_format.space_before = Pt(8)
        _add_run(references, "REFERENCES", theme)

    doc.add_page_break()


def _strip_duplicate_heading(draft, section):
    """Drop a leading draft line that merely repeats the section heading."""
    lines = draft.split("\n")
    # Skip leading blank lines before inspecting the first content line.
    first_index = 0
    while first_index < len(lines) and not lines[first_index].strip():
        first_index += 1
    if first_index >= len(lines):
        return draft

    def normalize(text):
        return re.sub(r"\s+", " ", re.sub(r"[*_`#]", "", text)).strip().lower()

    first = normalize(lines[first_index])
    expected = normalize(f"{section.get('id', '')} {section.get('heading', '')}")
    if first and (first == expected or first == normalize(section.get("heading", ""))):
        lines = lines[first_index + 1:]
    return "\n".join(lines).strip("\n")


VITAL_LABEL_RE = re.compile(
    r"temperature|pulse|respiration|blood pressure|spo₂|spo2|weight", re.I
)


def _vital_detail(label, value):
    """Vitals carry their unit in the label: "Temperature (°C)" -> "38.7°C"."""
    name = re.sub(r"\(.*\)", "", label).strip()
    unit_match = re.search(r"\(([^)]*)\)", label)
    unit = unit_match.group(1).strip() if unit_match else ""
    detail = value.strip()
    if unit and unit not in detail:
        detail = f"{detail}{unit}" if unit == "%" else f"{detail} {unit}"
    return f"{name} {detail}"
# Long student-written values are already prose — keep them as clean paragraphs.
PROSE_LENGTH = 60


def _prose_date(value):
    """'2026-08-13' -> '13th August, 2026'; anything else passes through."""
    match = re.fullmatch(r"(\d{4})-(\d{2})-(\d{2})", value.strip())
    if not match:
        return value
    year, month, day = match.groups()
    month_name = MONTH_NAMES.get(int(month))
    if month_name is None:
        return value
    day_number = int(day)
    suffix = (
        "th"
        if 11 <= day_number % 100 <= 13
        else {1: "st", 2: "nd", 3: "rd"}.get(day_number % 10, "th")
    )
    return f"{day_number}{suffix} {month_name}, {year}"


def _join_with_and(items):
    if len(items) <= 1:
        return "".join(items)
    return ", ".join(items[:-1]) + " and " + items[-1]


def _particulars_prose(parts):
    """1.1 Patient's Particulars -> one flowing biography paragraph instead of
    one mechanical sentence per field (mirrors particularsProse in App.tsx).

    Returns a list of (label_or_None, text) tuples."""

    def find(pattern):
        for part in parts:
            if pattern.search(part["label"]):
                return part["value"]
        return ""

    name = find(re.compile(r"name\s*/\s*initials|initials", re.I))
    age = find(re.compile(r"^age$", re.I))
    dob = find(re.compile(r"date of birth", re.I))
    sex = find(re.compile(r"^sex$", re.I))
    ethnicity = find(re.compile(r"ethnicity|tribe", re.I))
    religion = find(re.compile(r"religion", re.I))
    marital = find(re.compile(r"marital status", re.I))
    occupation = find(re.compile(r"occupation", re.I))
    address = find(re.compile(r"address|residence", re.I))
    ward = find(re.compile(r"ward|unit", re.I))
    # "Hospital number" must not be mistaken for the facility field.
    facility = find(re.compile(r"facility|hospital(?!\s*number)", re.I))
    admission_date = find(re.compile(r"date.*admission|admission.*date", re.I))
    diagnosis = find(re.compile(r"diagnosis", re.I))
    informant = find(re.compile(r"informant", re.I))

    prose = []

    identity = []
    if sex:
        identity.append(sex.lower())
    age_match = re.match(r"^\d+", age.strip())
    if age_match:
        identity.append(f"aged {age_match.group(0)} years")
    elif dob:
        identity.append(f"born on {_prose_date(dob)}")
    if ethnicity:
        identity.append(f"of the {ethnicity.strip()} tribe")
    if religion:
        rel = religion.strip()
        if re.search(r"christian|muslim|hindu|buddhist", rel, re.I):
            identity.append(f"a {rel}")
        else:
            identity.append(f"of the {rel} faith")
    if marital:
        identity.append(marital.lower())
    if occupation:
        identity.append(f"a {occupation.strip().lower()} by occupation")
    if identity:
        subject = f"The patient, {name.strip()}," if name else "The patient"
        residence = f", residing at {address.strip()}" if address else ""
        prose.append((None, f"{subject} is {_join_with_and(identity)}{residence}."))

    admission = []
    if ward:
        admission.append(f"to the {ward.strip().lower()}")
    if facility:
        admission.append(f"at {facility.strip()}")
    if admission_date:
        admission.append(f"on {_prose_date(admission_date.strip())}")
    if admission:
        if sex and sex.lower() == "female":
            pronoun = "she"
        elif sex:
            pronoun = "he"
        else:
            pronoun = "the patient"
        diagnosis_clause = (
            f" with a diagnosis of {diagnosis.strip()}" if diagnosis else ""
        )
        prose.append(
            (
                None,
                f"{pronoun.capitalize()} was admitted {', '.join(admission)}{diagnosis_clause}.",
            )
        )
    elif diagnosis:
        prose.append((None, f"The admission diagnosis was {diagnosis.strip()}."))

    if informant:
        cleaned = re.sub(r"\s*[-–—]\s*.*$", "", informant.strip()).strip()
        if re.match(r"^him", cleaned, re.I):
            phrase = "the patient himself"
        elif re.match(r"^her", cleaned, re.I):
            phrase = "the patient herself"
        else:
            phrase = cleaned
        reliability = (
            "; the information was deemed reliable"
            if re.search(r"reliab", informant, re.I)
            else ""
        )
        prose.append((None, f"The informant was {phrase}{reliability}."))

    cluster_re = re.compile(
        r"name\s*/\s*initials|initials|^age$|^sex$|ethnicity|tribe|religion|marital "
        r"status|occupation|address|residence|date of birth|ward|unit|facility|"
        r"hospital(?!\s*number)|admission|diagnosis|informant",
        re.I,
    )
    for part in parts:
        if not cluster_re.search(part["label"]):
            prose.append((part["label"], part["value"]))
    return prose


def _generic_prose(parts):
    """All other sections: long student-written values stay as clean paragraphs,
    short facts get a bold-label lead-in, and vitals group into one factual
    list (mirrors genericProse in App.tsx). Returns (label_or_None, text)."""
    prose = []

    vitals = [p for p in parts if VITAL_LABEL_RE.search(p["label"])]
    if vitals:
        details = [_vital_detail(p["label"], p["value"]) for p in vitals]
        text = ", ".join(details)
        prose.append((None, text[0].upper() + text[1:] + "."))
    vital_labels = {p["label"] for p in vitals}

    for part in parts:
        if part["label"] in vital_labels:
            continue
        value = part["value"]
        if len(value) >= PROSE_LENGTH:
            prose.append((None, value))
        else:
            prose.append((part["label"], value))
    return prose


def _field_prose(fields, section_id=""):
    """Compose a section's collected field data into professional prose for the
    Word export (mirrors fieldsToProse in App.tsx). Section 1.1 is composed
    into a flowing biography; every other section renders the student's own
    values as paragraphs with bold-label facts for short entries."""
    parts = [
        {"label": f.get("label") or "", "value": (f.get("value") or "").strip()}
        for f in fields
        if (f.get("value") or "").strip()
    ]
    if section_id == "1.1":
        return _particulars_prose(parts)
    return _generic_prose(parts)


def _render_section(doc, section, theme):
    section_heading = doc.add_paragraph(style=STYLE_HEADING_2)
    _add_run(section_heading, f"{section.get('id', '')} {section.get('heading', '')}".strip(), theme)

    draft = _strip_duplicate_heading((section.get("draft") or "").strip(), section)
    fields = [f for f in section.get("fields") or [] if (f.get("value") or "").strip()]
    rows = section.get("rows")
    row_data = rows.get("data") if rows else None
    # An AI-drafted table (e.g. the drugs table in 2.2) IS the section's
    # deliverable and replaces the structured scaffold — UNLESS it omits rows
    # the student entered (models occasionally drop items). In that case the
    # structured table is kept as a safety net so no collected data silently
    # disappears from the export.
    draft_table_rows = _draft_table_row_count(draft)
    student_row_count = len(row_data) if row_data else 0
    draft_has_table = draft_table_rows > 0
    draft_covers_rows = draft_table_rows >= student_row_count

    if draft:
        # A drafted paragraph is the section's content — the fields are NOT
        # dumped underneath it (the draft already narrates that data). This
        # also applies to mixed sections like 5.1: the draft paragraph
        # replaces the field list, while any structured rows still render.
        _add_draft(doc, draft, theme)
    elif fields:
        for label, text in _field_prose(fields, section.get("id", "")):
            p = doc.add_paragraph(style=STYLE_BODY)
            if label:
                _add_run(p, f"{label}: ", theme, bold=True)
            _add_inline_text(p, _format_ordinal_dates(text), theme)
    if row_data and (not draft_has_table or not draft_covers_rows):
        _add_rows_table(doc, rows, theme)
    if not draft and not fields and not row_data:
        p = doc.add_paragraph(style=STYLE_BODY)
        _add_run(p, "Not completed.", theme, italic=True)


def _add_chapter(doc, chapters, chapter_index, chapter, theme, include_intro=True):
    chapter_heading = doc.add_paragraph(style=STYLE_HEADING_1)
    _add_run(chapter_heading, _chapter_heading(chapters, chapter_index, chapter), theme)

    intro = (chapter.get("intro") or "").strip()
    if include_intro and intro:
        # Match the sample care studies: a short opening paragraph right under
        # the chapter heading, rendered like a drafted section's prose.
        _add_draft(doc, intro, theme)

    for section in chapter.get("sections", []):
        _render_section(doc, section, theme)


def _collect_references(chapters):
    """Distinct reference labels cited across the given chapters, in order.

    Includes chapter introductions' sources first (they appear earliest in the
    document), then every section's sources."""
    seen = set()
    refs = []

    def add_ref(ref):
        label = (ref.get("label") or "").strip()
        if not label or label in seen:
            return
        seen.add(label)
        refs.append(label)

    for chapter in chapters:
        for ref in chapter.get("introReferences") or []:
            add_ref(ref)
        for section in chapter.get("sections", []):
            for ref in section.get("references") or []:
                add_ref(ref)
    return refs


def _has_bibliography(chapters):
    """True when the student's Bibliography section (6.3) has entries. The
    curated bibliography then replaces the auto-generated REFERENCES page, so
    the document never carries two reference lists."""
    for chapter in chapters:
        for section in chapter.get("sections", []):
            if section.get("id") == "6.3":
                rows = section.get("rows") or {}
                return bool(rows.get("data"))
    return False


def _add_references(doc, chapters, theme, page_break_before=False):
    """End-of-document REFERENCES list, matching the sample care studies."""
    refs = _collect_references(chapters)
    if not refs:
        return
    if page_break_before:
        doc.add_page_break()
    heading = doc.add_paragraph(style=STYLE_TOC_TITLE)
    # Keep the original (non-heading) spacing of this centered heading.
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(6)
    _add_run(heading, "REFERENCES", theme)
    for index, label in enumerate(refs, start=1):
        p = doc.add_paragraph(style=STYLE_BODY)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)  # hanging indent
        _add_run(p, f"{index}. ", theme)
        _add_inline_text(p, _format_ordinal_dates(label), theme)


def build_docx(payload):
    theme = Theme.from_dict(payload.get("theme"))
    doc = Document()
    _setup_styles(doc, theme)
    _setup_lists(doc, theme)
    _add_page_number(doc)

    title = payload.get("title") or {}
    chapters = payload.get("chapters") or []
    scope = payload.get("scope") or {}
    scope_type = scope.get("type") or "full"
    # Unknown/invalid scope types render the full study — degrade gracefully
    # rather than leaving chapter/section variables unbound below.
    if scope_type not in ("chapter", "section"):
        scope_type = "full"

    if scope_type in ("chapter", "section"):
        chapter_index = int(scope.get("chapterIndex") or 0)
        chapter = chapters[chapter_index] if 0 <= chapter_index < len(chapters) else None
        if chapter is None:
            scope_type = "full"  # fall back rather than emit an empty document
        elif scope_type == "section":
            section_index = int(scope.get("sectionIndex") or 0)
            sections = chapter.get("sections") or []
            if 0 <= section_index < len(sections):
                # A single-section export renders just that section — the whole
                # chapter's introduction would be context the student didn't ask for.
                chapter = {**chapter, "sections": [sections[section_index]], "intro": ""}
            else:
                scope_type = "full"

    if scope_type == "full":
        _add_title_page(doc, title, theme)
        _add_toc(doc, chapters, theme)
        for chapter_index, chapter in enumerate(chapters):
            _add_chapter(doc, chapters, chapter_index, chapter, theme)
        if not _has_bibliography(chapters):
            _add_references(doc, chapters, theme, page_break_before=True)
    else:
        # Chapter/section exports begin directly with the chapter heading — no
        # title page, TOC, or header block; the full study alone carries those.
        _add_chapter(doc, chapters, chapter_index, chapter, theme)
        if not _has_bibliography([chapter]):
            _add_references(doc, [chapter], theme)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--file", help="Path to a JSON file (defaults to stdin)")
    args = parser.parse_args()

    if args.file:
        with open(args.file, "r", encoding="utf-8") as f:
            payload = json.load(f)
    else:
        # Read stdin as UTF-8 bytes: Windows pipes default to cp1252, which
        # cannot decode every character a drafted section may contain.
        payload = json.loads(sys.stdin.buffer.read().decode("utf-8"))

    data = build_docx(payload)
    sys.stdout.buffer.write(data)
    sys.stdout.buffer.flush()


if __name__ == "__main__":
    main()
